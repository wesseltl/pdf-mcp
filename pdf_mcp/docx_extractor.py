"""Extract text and tables from Word .docx files.

Word tables are already structured, so extraction is deterministic: cells are read from the document
model rather than inferred from page geometry. The same table assessment used for PDFs is applied so
callers get consistent reliability fields.
"""
from __future__ import annotations

import csv
import io

from docx import Document

from pdf_mcp.extractor import _assess, _check_path


def extract_docx_text(path: str) -> dict:
    """Extract non-empty paragraph text from a .docx file."""
    doc = Document(_check_path(path))
    paragraphs = [
        {"index": i + 1, "text": paragraph.text.strip()}
        for i, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip()
    ]
    return {
        "paragraphs": paragraphs,
        "text": "\n".join(p["text"] for p in paragraphs),
        "n_paragraphs": len(paragraphs),
    }


def extract_docx_tables(path: str) -> dict:
    """Extract Word tables as rows of cell strings."""
    doc = Document(_check_path(path))
    tables = []
    for i, table in enumerate(doc.tables):
        rows = []
        has_merged_cells = False
        for row in table.rows:
            cells = list(row.cells)
            if len({id(cell._tc) for cell in cells}) < len(cells):
                has_merged_cells = True
            rows.append([cell.text.strip() for cell in cells])
        assessment = _assess(rows)
        if has_merged_cells:
            assessment["warnings"] = list(assessment.get("warnings", [])) + [
                "merged cells detected (Word exposes merged cells as repeated values)"
            ]
            assessment["looks_clean"] = False
        tables.append({
            "index": i,
            "rows": rows,
            "n_rows": len(rows),
            "has_merged_cells": has_merged_cells,
            **assessment,
        })
    return {"tables": tables, "n_tables": len(tables)}


def docx_table_to_csv(path: str, index: int = 0) -> str:
    """Return one extracted Word table (default the first) as CSV text."""
    tables = extract_docx_tables(path)["tables"]
    if not tables:
        return ""
    rows = tables[index]["rows"]
    buf = io.StringIO()
    csv.writer(buf).writerows(rows)
    return buf.getvalue()
