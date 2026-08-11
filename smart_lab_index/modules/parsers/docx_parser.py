"""Word DOCX parser module."""

from __future__ import annotations

from typing import BinaryIO

from docx import Document

from smart_lab_index.core.domain import (
    DocumentContent,
    DocumentSource,
    TableContent,
    TextBlock,
)
from smart_lab_index.core.modules import (
    FileAccess,
    ModuleCapability,
    ModuleManifest,
    ModuleType,
    NetworkAccess,
    ParserModule,
)
from smart_lab_index.modules.parsers.common import (
    MAX_TEXT_BLOCKS,
    DocumentParseError,
    assess_table,
    preflight_office_archive,
    provenance,
    require_table_budget,
    require_text_budget,
    source_extension,
)

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class DocxParser(ParserModule):
    manifest = ModuleManifest(
        module_id="parser.docx",
        name="DOCX Parser",
        version="0.2.0",
        module_type=ModuleType.PARSER,
        description="Extracts paragraphs and tables from Word document streams.",
        capabilities=(ModuleCapability("parser.document", "1.0.0"),),
        network_access=NetworkAccess.NONE,
        file_access=FileAccess.NONE,
    )

    def supports(self, source: DocumentSource) -> bool:
        return source_extension(source) == ".docx" or source.content_type == DOCX_CONTENT_TYPE

    def parse(self, source: DocumentSource, content: BinaryIO) -> DocumentContent:
        try:
            preflight_office_archive(content, "DOCX document")
            document = Document(content)
            blocks = []
            text_characters = 0
            for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
                text = paragraph.text.strip()
                if text:
                    if len(blocks) >= MAX_TEXT_BLOCKS:
                        raise DocumentParseError("DOCX document exceeds the text block limit")
                    text_characters += len(text)
                    require_text_budget(text_characters, "DOCX document")
                    blocks.append(TextBlock(
                        index=len(blocks),
                        kind="paragraph",
                        text=text,
                        provenance=provenance(
                            source,
                            {"paragraph": paragraph_index},
                            excerpt=text[:240],
                        ),
                    ))
            tables = []
            table_rows = 0
            table_cells = 0
            for table_index, table in enumerate(document.tables):
                rows = []
                cell_references = []
                has_merged_cells = False
                for row_index, row in enumerate(table.rows, start=1):
                    cells = list(row.cells)
                    table_rows += 1
                    table_cells += len(cells)
                    require_table_budget(table_rows, table_cells, "DOCX document")
                    if len({id(cell._tc) for cell in cells}) < len(cells):
                        has_merged_cells = True
                    rows.append(tuple(cell.text.strip() for cell in cells))
                    cell_references.append(tuple(
                        provenance(source, {
                            "table": table_index,
                            "row": row_index,
                            "column": column_index,
                        })
                        for column_index in range(1, len(cells) + 1)
                    ))
                assessment = assess_table(rows)
                warnings = list(assessment["warnings"])
                if has_merged_cells:
                    warnings.append("merged Word cells are exposed as repeated values")
                tables.append(TableContent(
                    index=table_index,
                    name=f"Table {table_index + 1}",
                    rows=tuple(rows),
                    cell_provenance=tuple(cell_references),
                    metadata={
                        **assessment,
                        "has_merged_cells": has_merged_cells,
                        "warnings": warnings,
                    },
                ))
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError(f"DOCX parsing failed: {exc}") from exc
        warnings = []
        if not blocks:
            warnings.append("no Word paragraph text detected")
        return DocumentContent(
            source_external_id=source.external_id,
            content_type=source.content_type,
            parser_module_id=self.manifest.module_id,
            parser_version=self.manifest.version,
            text_blocks=tuple(blocks),
            tables=tuple(tables),
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            },
            warnings=tuple(warnings),
        )
