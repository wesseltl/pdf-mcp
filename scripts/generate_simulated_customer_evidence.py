"""Generate a deterministic, explicitly fictional customer evaluation pack."""
from __future__ import annotations

import argparse
import json
import sys
import zipfile
from collections import Counter
from datetime import datetime
from decimal import Decimal
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from pdf_mcp.evaluator import evaluate_manifest  # noqa: E402


DEFAULT_OUTPUT = ROOT / "evaluations" / "simulated-customer"
FIXED_DOCUMENT_TIME = datetime(2026, 1, 1, 0, 0, 0)
FIXED_ZIP_TIME = (2026, 1, 1, 0, 0, 0)


PROFILE = {
    "$schema": "../../profile.schema.json",
    "profile_schema_version": "1.0",
    "id": "northstar-water-results-v1",
    "version": "1.0.0",
    "description": (
        "Fictional Northstar Water result rows used only for simulated customer evaluation."
    ),
    "table": {
        "header_search_rows": 3,
        "minimum_header_match": 1.0,
        "allow_extra_columns": True,
        "min_records": 1,
        "unique_by": ["sample_id", "analyte"],
        "columns": [
            {
                "name": "sample_id",
                "aliases": ["Sample ID", "Sample", "Sample No.", "Sample #", "Specimen"],
                "type": "string",
                "required": True,
                "allow_blank": False,
                "pattern": "^NS-[0-9]{4}$",
            },
            {
                "name": "analyte",
                "aliases": ["Analyte", "Test", "Parameter"],
                "type": "string",
                "required": True,
                "allow_blank": False,
            },
            {
                "name": "result",
                "aliases": ["Result", "Value", "Measured Value", "Measurement"],
                "type": "decimal",
                "required": True,
                "allow_blank": False,
                "minimum": "0",
                "maximum": "100000",
            },
            {
                "name": "unit",
                "aliases": ["Unit", "Units", "UOM"],
                "type": "string",
                "required": True,
                "allow_blank": False,
            },
        ],
    },
}


def _json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2) + "\n", encoding="utf-8")


def _pdf_table(rows: list[list[str]], span_first_row: bool = False):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    table = Table(rows)
    style = [
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    if span_first_row:
        style.append(("SPAN", (0, 0), (-1, 0)))
    table.setStyle(TableStyle(style))
    return table


def _invariant_canvas(filename, *args, **kwargs):
    from reportlab.pdfgen import canvas

    kwargs["invariant"] = 1
    return canvas.Canvas(filename, *args, **kwargs)


def write_pdf(path: Path, pages: list[list[tuple[list[list[str]], bool]]]) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import PageBreak, SimpleDocTemplate, Spacer

    story = []
    for page_index, tables in enumerate(pages):
        for table_index, (rows, span_first_row) in enumerate(tables):
            if table_index:
                story.append(Spacer(1, 28))
            story.append(_pdf_table(rows, span_first_row))
        if page_index < len(pages) - 1:
            story.append(PageBreak())
    path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(path), pagesize=A4, title="Simulated Northstar Water Results"
    )
    document.build(story, canvasmaker=_invariant_canvas)


def write_text_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(str(path), pagesize=A4, title="Simulated cover note")
    document.build(
        [Paragraph("No analytical result table is present.", getSampleStyleSheet()["BodyText"])],
        canvasmaker=_invariant_canvas,
    )


def _canonicalize_docx(raw_path: Path, output_path: Path) -> None:
    """Normalize ZIP metadata so generated DOCX hashes are reproducible."""
    with zipfile.ZipFile(raw_path, "r") as source, zipfile.ZipFile(
        output_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as destination:
        for name in sorted(source.namelist()):
            original = source.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            destination.writestr(info, source.read(name))
    raw_path.unlink()


def write_docx(
    path: Path,
    tables: list[list[list[str]]],
    merged_title: bool = False,
) -> None:
    from docx import Document

    document = Document()
    document.core_properties.title = "Simulated Northstar Water Results"
    document.core_properties.author = "pdf-mcp synthetic fixture generator"
    document.core_properties.created = FIXED_DOCUMENT_TIME
    document.core_properties.modified = FIXED_DOCUMENT_TIME
    for table_index, rows in enumerate(tables):
        width = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=width)
        for row_index, row in enumerate(rows):
            for column_index, value in enumerate(row):
                table.rows[row_index].cells[column_index].text = value
        if merged_title and table_index == 0:
            table.rows[0].cells[0].merge(table.rows[0].cells[-1])
        if table_index < len(tables) - 1:
            document.add_paragraph("")
    path.parent.mkdir(parents=True, exist_ok=True)
    raw_path = path.with_name(f".{path.name}.raw")
    document.save(raw_path)
    _canonicalize_docx(raw_path, path)


def expected(rows: list[list[str]]) -> list[dict]:
    records = []
    for sample_id, analyte, result, unit, *_extra in rows:
        records.append({
            "sample_id": sample_id,
            "analyte": analyte,
            "result": format(Decimal(result), "f"),
            "unit": unit or None,
        })
    return records


def case(case_id: str, filename: str, decision: str, records: list[dict]) -> dict:
    return {
        "id": case_id,
        "document": f"fixtures/{filename}",
        "expected_decision": decision,
        "expected_records": records,
    }


def generate(output: Path) -> tuple[Path, Path]:
    fixtures = output / "fixtures"
    fixtures.mkdir(parents=True, exist_ok=True)
    _json(output / "profile.json", PROFILE)

    canonical_header = ["Sample ID", "Analyte", "Result", "Unit"]

    dev_standard = [
        ["NS-1001", "pH", "7.20", "pH units"],
        ["NS-1001", "Chloride", "125", "mg/L"],
        ["NS-1002", "Turbidity", "0.45", "NTU"],
    ]
    write_pdf(fixtures / "dev-01-standard.pdf", [[([canonical_header, *dev_standard], False)]])

    dev_aliases = [
        ["NS-1010", "Nitrate", "4.5", "mg/L"],
        ["NS-1010", "Conductivity", "415", "uS/cm"],
    ]
    write_pdf(
        fixtures / "dev-02-aliases.pdf",
        [[([ ["Sample", "Test", "Value", "UOM"], *dev_aliases], False)]],
    )

    dev_title = [
        ["NS-1020", "Iron", "0.12", "mg/L"],
        ["NS-1020", "Manganese", "0.03", "mg/L"],
    ]
    write_pdf(
        fixtures / "dev-03-title-row.pdf",
        [[([ ["Northstar Water Results", "", "", ""], canonical_header, *dev_title], True)]],
    )

    dev_extra = [
        ["NS-1030", "Hardness", "180", "mg/L", "SM 2340 C"],
        ["NS-1030", "Alkalinity", "95", "mg/L", "SM 2320 B"],
    ]
    write_pdf(
        fixtures / "dev-04-extra-column.pdf",
        [[([ [*canonical_header, "Method"], *dev_extra], False)]],
    )

    dev_docx = [
        ["NS-1040", "Calcium", "42.5", "mg/L"],
        ["NS-1040", "Magnesium", "11.2", "mg/L"],
    ]
    write_docx(fixtures / "dev-05-standard.docx", [[canonical_header, *dev_docx]])

    dev_table_one = [["NS-1050", "pH", "7.1", "pH units"]]
    dev_table_two = [["NS-1051", "Chloride", "88", "mg/L"]]
    write_pdf(
        fixtures / "dev-06-two-tables.pdf",
        [[
            ([canonical_header, *dev_table_one], False),
            ([canonical_header, *dev_table_two], False),
        ]],
    )

    dev_page_one = [
        ["NS-1060", "pH", "7.0", "pH units"],
        ["NS-1060", "Chloride", "101", "mg/L"],
    ]
    dev_page_two = [
        ["NS-1061", "Turbidity", "0.8", "NTU"],
        ["NS-1061", "Nitrate", "2.2", "mg/L"],
    ]
    write_pdf(
        fixtures / "dev-07-headerless-continuation.pdf",
        [
            [([canonical_header, *dev_page_one], False)],
            [(dev_page_two, False)],
        ],
    )

    dev_bad_id = [["SAMPLE-1070", "pH", "7.3", "pH units"]]
    write_pdf(
        fixtures / "dev-08-invalid-sample-id.pdf",
        [[([canonical_header, *dev_bad_id], False)]],
    )

    dev_negative = [["NS-1080", "Turbidity", "-0.2", "NTU"]]
    write_pdf(
        fixtures / "dev-09-negative-result.pdf",
        [[([canonical_header, *dev_negative], False)]],
    )

    dev_merged = [["Northstar Results", "", "", ""], canonical_header,
                  ["NS-1090", "Nitrate", "3.4", "mg/L"]]
    write_docx(fixtures / "dev-10-merged-title.docx", [dev_merged], merged_title=True)

    write_pdf(
        fixtures / "dev-11-unrelated.pdf",
        [[([["Date", "Operator"], ["2026-01-01", "A. Example"]], False)]],
    )

    write_pdf(
        fixtures / "dev-12-partial-header.pdf",
        [[([ ["Sample ID", "Analyte"], ["NS-1120", "pH"]], False)]],
    )

    development_cases = [
        case("dev-01-standard", "dev-01-standard.pdf", "accepted", expected(dev_standard)),
        case("dev-02-aliases", "dev-02-aliases.pdf", "accepted", expected(dev_aliases)),
        case("dev-03-title-row", "dev-03-title-row.pdf", "accepted", expected(dev_title)),
        case("dev-04-extra-column", "dev-04-extra-column.pdf", "accepted", expected(dev_extra)),
        case("dev-05-standard-docx", "dev-05-standard.docx", "accepted", expected(dev_docx)),
        case(
            "dev-06-two-tables", "dev-06-two-tables.pdf", "accepted",
            expected([*dev_table_one, *dev_table_two]),
        ),
        case(
            "dev-07-headerless-continuation", "dev-07-headerless-continuation.pdf",
            "needs_review", expected([*dev_page_one, *dev_page_two]),
        ),
        case(
            "dev-08-invalid-sample-id", "dev-08-invalid-sample-id.pdf", "needs_review",
            expected(dev_bad_id),
        ),
        case(
            "dev-09-negative-result", "dev-09-negative-result.pdf", "needs_review",
            expected(dev_negative),
        ),
        case(
            "dev-10-merged-title", "dev-10-merged-title.docx", "needs_review",
            expected([dev_merged[-1]]),
        ),
        case("dev-11-unrelated", "dev-11-unrelated.pdf", "rejected", []),
        case("dev-12-partial-header", "dev-12-partial-header.pdf", "rejected", []),
    ]

    hold_alias = [
        ["NS-2001", "Sulfate", "31", "mg/L"],
        ["NS-2001", "Fluoride", "0.72", "mg/L"],
    ]
    write_pdf(
        fixtures / "holdout-01-punctuation-aliases.pdf",
        [[([ ["SAMPLE #", "PARAMETER", "MEASURED VALUE", "UNITS"], *hold_alias], False)]],
    )

    hold_docx = [
        ["NS-2002", "Copper", "0.08", "mg/L"],
        ["NS-2002", "Zinc", "0.15", "mg/L"],
    ]
    write_docx(
        fixtures / "holdout-02-docx-aliases.docx",
        [[["Specimen", "Parameter", "Measurement", "Units"], *hold_docx]],
    )

    hold_page_one = [
        ["NS-2003", "pH", "6.95", "pH units"],
        ["NS-2003", "Turbidity", "0.35", "NTU"],
    ]
    hold_page_two = [
        ["NS-2004", "Chloride", "76", "mg/L"],
        ["NS-2004", "Nitrate", "1.8", "mg/L"],
    ]
    write_pdf(
        fixtures / "holdout-03-repeated-header.pdf",
        [
            [([canonical_header, *hold_page_one], False)],
            [([canonical_header, *hold_page_two], False)],
        ],
    )

    hold_blank_unit = [["NS-2005", "Conductivity", "390", ""]]
    write_pdf(
        fixtures / "holdout-04-blank-unit.pdf",
        [[([canonical_header, *hold_blank_unit], False)]],
    )

    hold_duplicate = [
        ["NS-2006", "Iron", "0.10", "mg/L"],
        ["NS-2006", "Iron", "0.10", "mg/L"],
    ]
    write_pdf(
        fixtures / "holdout-05-duplicate-key.pdf",
        [[([canonical_header, *hold_duplicate], False)]],
    )

    write_text_pdf(fixtures / "holdout-06-text-only.pdf")

    holdout_cases = [
        case(
            "holdout-01-punctuation-aliases", "holdout-01-punctuation-aliases.pdf",
            "accepted", expected(hold_alias),
        ),
        case(
            "holdout-02-docx-aliases", "holdout-02-docx-aliases.docx", "accepted",
            expected(hold_docx),
        ),
        case(
            "holdout-03-repeated-header", "holdout-03-repeated-header.pdf", "accepted",
            expected([*hold_page_one, *hold_page_two]),
        ),
        case(
            "holdout-04-blank-unit", "holdout-04-blank-unit.pdf", "needs_review",
            expected(hold_blank_unit),
        ),
        case(
            "holdout-05-duplicate-key", "holdout-05-duplicate-key.pdf", "needs_review",
            expected(hold_duplicate),
        ),
        case("holdout-06-text-only", "holdout-06-text-only.pdf", "rejected", []),
    ]

    common = {
        "evaluation_schema_version": "1.0",
        "evidence_label": "simulated_fictional_customer",
        "profile": "profile.json",
        "minimums": {
            "field_precision": 1.0,
            "field_recall": 1.0,
            "field_f1": 1.0,
            "exact_record_rate": 1.0,
            "decision_accuracy": 1.0,
        },
    }
    development_manifest = output / "development.json"
    holdout_manifest = output / "holdout.json"
    _json(development_manifest, {**common, "cases": development_cases})
    _json(holdout_manifest, {**common, "cases": holdout_cases})

    development_report = evaluate_manifest(str(development_manifest))
    holdout_report = evaluate_manifest(str(holdout_manifest))
    _json(output / "development-report.json", development_report)
    _json(output / "holdout-report.json", holdout_report)
    decisions = Counter(
        item["expected_decision"] for item in [*development_cases, *holdout_cases]
    )
    _json(output / "summary.json", {
        "evidence_label": "simulated_fictional_customer",
        "fictional_customer": "Northstar Water Operations",
        "not_real_customer_evidence": True,
        "profile": development_report["profile"],
        "documents": len(development_cases) + len(holdout_cases),
        "development_documents": len(development_cases),
        "holdout_documents": len(holdout_cases),
        "expected_decisions": {
            decision: decisions[decision]
            for decision in ("accepted", "needs_review", "rejected")
        },
        "development_metrics": development_report["metrics"],
        "holdout_metrics": holdout_report["metrics"],
        "limitations": [
            "All documents and expected values were generated by the product author.",
            "The fixtures do not demonstrate customer demand, retention, or willingness to pay.",
            "The measured scores do not transfer to unseen real document families.",
        ],
    })
    return development_manifest, holdout_manifest


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Generate explicitly fictional pdf-mcp customer evidence fixtures."
    )
    parser.add_argument(
        "--output-dir", type=Path, default=DEFAULT_OUTPUT,
        help=f"Output directory (default: {DEFAULT_OUTPUT})",
    )
    args = parser.parse_args(argv)
    development, holdout = generate(args.output_dir.resolve())
    print(f"Generated simulated development evidence: {development}")
    print(f"Generated simulated holdout evidence: {holdout}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
