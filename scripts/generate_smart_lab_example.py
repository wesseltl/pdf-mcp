"""Generate a small, entirely synthetic Smart Lab Index demonstration."""

from __future__ import annotations

import argparse
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import Workbook

_FIXED_DATETIME = datetime(2026, 1, 1, 0, 0, 0)  # noqa: DTZ001 - Office uses naive time
_FIXED_ZIP_TIME = (2026, 1, 1, 0, 0, 0)
_FIXED_W3CDTF = "2026-01-01T00:00:00Z"
_CORE_NAMESPACES = {
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
    "xsi": "http://www.w3.org/2001/XMLSchema-instance",
}
for _prefix, _uri in _CORE_NAMESPACES.items():
    ET.register_namespace(_prefix, _uri)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("examples/smart_lab_index/sample_lab"),
    )
    args = parser.parse_args()
    generate(args.output)
    print(args.output.resolve())
    return 0


def generate(output: Path) -> None:
    output.mkdir(parents=True, exist_ok=True)
    _workbook(
        output / "locations.xlsx",
        "Locations",
        [
            ["location_id", "name", "location_type"],
            ["Room A-101", "Room A-101", "ROOM"],
            ["Room A-102", "Room A-102", "ROOM"],
        ],
    )
    _workbook(
        output / "equipment.xlsx",
        "Assets",
        [
            ["asset_id", "name", "asset_type", "location"],
            ["Freezer-001", "Freezer-001", "FREEZER", "Room A-101"],
        ],
    )
    _workbook(
        output / "responsibilities.xlsx",
        "Responsibilities",
        [
            ["person", "asset_id", "relationship"],
            ["Alex Example", "Freezer-001", "responsible_for"],
        ],
    )
    sop_directory = output / "SOPs"
    sop_directory.mkdir(exist_ok=True)
    document = Document()
    document.core_properties.created = _FIXED_DATETIME
    document.core_properties.modified = _FIXED_DATETIME
    document.add_heading("Freezer procedure", level=1)
    document.add_paragraph("Freezer-001 located in Room A-102.")
    sop_path = sop_directory / "SOP_freezers.docx"
    document.save(sop_path)
    _normalize_office_zip(sop_path)


def _workbook(path: Path, sheet_name: str, rows: list[list[str]]) -> None:
    workbook = Workbook()
    workbook.properties.created = _FIXED_DATETIME
    workbook.properties.modified = _FIXED_DATETIME
    sheet = workbook.active
    sheet.title = sheet_name
    for row in rows:
        sheet.append(row)
    workbook.save(path)
    workbook.close()
    _normalize_office_zip(path)


def _normalize_office_zip(path: Path) -> None:
    replacement = path.with_suffix(f"{path.suffix}.tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        replacement,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as destination:
        for name in sorted(source.namelist()):
            information = zipfile.ZipInfo(name, _FIXED_ZIP_TIME)
            information.compress_type = zipfile.ZIP_DEFLATED
            information.create_system = 3
            information.external_attr = (0o755 if name.endswith("/") else 0o644) << 16
            destination.writestr(
                information,
                _normalize_office_member(name, source.read(name)),
            )
    replacement.replace(path)


def _normalize_office_member(name: str, content: bytes) -> bytes:
    if name != "docProps/core.xml":
        return content
    root = ET.fromstring(content)
    for field in ("created", "modified"):
        element = root.find(f"{{{_CORE_NAMESPACES['dcterms']}}}{field}")
        if element is not None:
            element.text = _FIXED_W3CDTF
    return ET.tostring(root, encoding="utf-8")


if __name__ == "__main__":
    raise SystemExit(main())
