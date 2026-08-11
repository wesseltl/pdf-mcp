"""Contract tests for Smart Lab Index document parser modules."""

from __future__ import annotations

import io
import tempfile
import unittest
from collections.abc import Callable
from pathlib import Path
from unittest import mock

from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from scripts.generate_smart_lab_example import generate
from smart_lab_index.core.domain import DocumentContent, DocumentSource
from smart_lab_index.core.modules import ParserModule
from smart_lab_index.modules.parsers import (
    CsvParser,
    DocxParser,
    PdfParser,
    TextParser,
    XlsxParser,
)
from smart_lab_index.modules.parsers.common import DocumentParseError

CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".csv": "text/csv",
    ".txt": "text/plain",
}


def make_source(extension: str) -> DocumentSource:
    name = f"sample{extension}"
    return DocumentSource(
        external_id=f"documents/{name}",
        source_id="source.synthetic",
        name=name,
        path=f"Lab Alpha/{name}",
        content_type=CONTENT_TYPES[extension],
        modified_at="2026-01-02T03:04:05+00:00",
        size_bytes=0,
        checksum="synthetic-checksum",
    )


def make_pdf() -> io.BytesIO:
    stream = io.BytesIO()
    document = SimpleDocTemplate(stream, pagesize=A4)
    styles = getSampleStyleSheet()
    table = Table([
        ["Asset", "Location"],
        ["Freezer-001", "Room A-101"],
    ])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)]))
    document.build([
        Paragraph("Lab Alpha equipment", styles["Title"]),
        Spacer(1, 12),
        table,
    ])
    return stream


def make_docx() -> io.BytesIO:
    stream = io.BytesIO()
    document = Document()
    document.add_paragraph("Freezer-001 is located in Room A-101.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Asset"
    table.cell(0, 1).text = "Owner"
    table.cell(1, 0).text = "Freezer-001"
    table.cell(1, 1).text = "Alex Example"
    document.save(stream)
    return stream


def make_xlsx() -> io.BytesIO:
    stream = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Assets"
    sheet.append(["Asset", "Location"])
    sheet.append(["Freezer-001", "Room A-101"])
    workbook.create_sheet("Empty")
    workbook.save(stream)
    workbook.close()
    return stream


def make_csv() -> io.BytesIO:
    return io.BytesIO(
        b"Asset;Location\r\nFreezer-001;Room A-101\r\n"
    )


def make_txt() -> io.BytesIO:
    return io.BytesIO(
        "\ufeffLab Alpha\n\nFreezer-001 located in Room A-101\n".encode("utf-8")
    )


class TestSmartLabParserContracts(unittest.TestCase):
    def assert_document_contract(
        self,
        parser: ParserModule,
        source: DocumentSource,
        content: DocumentContent,
    ) -> None:
        self.assertEqual(content.source_external_id, source.external_id)
        self.assertEqual(content.content_type, source.content_type)
        self.assertEqual(content.parser_module_id, parser.manifest.module_id)
        self.assertEqual(content.parser_version, parser.manifest.version)
        self.assertTrue(parser.supports(source))
        self.assertFalse(hasattr(source, "content_ref"))

    def test_pdf_normalizes_page_text_table_and_cell_provenance(self) -> None:
        parser = PdfParser()
        source = make_source(".pdf")

        result = parser.parse(source, make_pdf())

        self.assert_document_contract(parser, source, result)
        self.assertEqual(result.metadata["page_count"], 1)
        self.assertIn("Lab Alpha equipment", result.text_blocks[0].text)
        self.assertEqual(result.text_blocks[0].kind, "page_text")
        self.assertEqual(result.text_blocks[0].provenance.locator, {"page": 1})
        self.assertEqual(
            result.tables[0].rows,
            (("Asset", "Location"), ("Freezer-001", "Room A-101")),
        )
        cell = result.tables[0].cell_provenance[1][0]
        self.assertEqual(cell.source_external_id, source.external_id)
        self.assertEqual(
            {key: cell.locator[key] for key in ("page", "table", "row", "column")},
            {"page": 1, "table": 0, "row": 2, "column": 1},
        )
        self.assertEqual(len(cell.locator["bbox"]), 4)

    def test_docx_normalizes_paragraph_table_and_cell_provenance(self) -> None:
        parser = DocxParser()
        source = make_source(".docx")

        result = parser.parse(source, make_docx())

        self.assert_document_contract(parser, source, result)
        self.assertEqual(
            [block.text for block in result.text_blocks],
            ["Freezer-001 is located in Room A-101."],
        )
        self.assertEqual(result.text_blocks[0].provenance.locator, {"paragraph": 1})
        self.assertEqual(
            result.tables[0].rows,
            (("Asset", "Owner"), ("Freezer-001", "Alex Example")),
        )
        cell = result.tables[0].cell_provenance[1][1]
        self.assertEqual(cell.source_external_id, source.external_id)
        self.assertEqual(cell.locator, {"table": 0, "row": 2, "column": 2})

    def test_xlsx_normalizes_nonempty_sheets_and_cell_provenance(self) -> None:
        parser = XlsxParser()
        source = make_source(".xlsx")

        result = parser.parse(source, make_xlsx())

        self.assert_document_contract(parser, source, result)
        self.assertEqual(result.metadata["sheet_names"], ["Assets", "Empty"])
        self.assertEqual(len(result.tables), 1)
        self.assertEqual(result.tables[0].name, "Assets")
        self.assertEqual(
            result.tables[0].rows,
            (("Asset", "Location"), ("Freezer-001", "Room A-101")),
        )
        cell = result.tables[0].cell_provenance[1][0]
        self.assertEqual(cell.source_external_id, source.external_id)
        self.assertEqual(
            cell.locator,
            {"sheet": "Assets", "row": 2, "column": 1, "cell": "A2"},
        )

    def test_csv_normalizes_detected_delimiter_and_cell_provenance(self) -> None:
        parser = CsvParser()
        source = make_source(".csv")

        result = parser.parse(source, make_csv())

        self.assert_document_contract(parser, source, result)
        self.assertEqual(result.tables[0].metadata["dialect_delimiter"], ";")
        self.assertEqual(
            result.tables[0].rows,
            (("Asset", "Location"), ("Freezer-001", "Room A-101")),
        )
        cell = result.tables[0].cell_provenance[1][1]
        self.assertEqual(cell.source_external_id, source.external_id)
        self.assertEqual(cell.locator, {"row": 2, "column": 2})

    def test_txt_normalizes_nonempty_lines_and_original_line_numbers(self) -> None:
        parser = TextParser()
        source = make_source(".txt")

        result = parser.parse(source, make_txt())

        self.assert_document_contract(parser, source, result)
        self.assertEqual(
            [block.text for block in result.text_blocks],
            ["Lab Alpha", "Freezer-001 located in Room A-101"],
        )
        self.assertEqual([block.index for block in result.text_blocks], [0, 1])
        self.assertEqual(
            [block.provenance.locator for block in result.text_blocks],
            [{"line": 1}, {"line": 3}],
        )
        self.assertEqual(result.metadata, {"encoding": "utf-8", "line_count": 3})

    def test_malformed_streams_raise_normalized_error_and_parser_recovers(self) -> None:
        cases: tuple[
            tuple[ParserModule, str, bytes, Callable[[], io.BytesIO]], ...
        ] = (
            (PdfParser(), ".pdf", b"not a PDF", make_pdf),
            (DocxParser(), ".docx", b"not a DOCX", make_docx),
            (XlsxParser(), ".xlsx", b"not an XLSX", make_xlsx),
            (CsvParser(), ".csv", b"\xff", make_csv),
            (TextParser(), ".txt", b"\xff", make_txt),
        )

        for parser, extension, malformed, valid_factory in cases:
            source = make_source(extension)
            with self.subTest(parser=parser.manifest.module_id):
                with self.assertRaises(DocumentParseError):
                    parser.parse(source, io.BytesIO(malformed))

                recovered = parser.parse(source, valid_factory())
                self.assert_document_contract(parser, source, recovered)
                self.assertTrue(recovered.text_blocks or recovered.tables)

    def test_parser_resource_budgets_reject_excessive_inputs(self) -> None:
        with (
            mock.patch("smart_lab_index.modules.parsers.common.MAX_TEXT_CHARS", 8),
            self.assertRaisesRegex(DocumentParseError, "text limit"),
        ):
            TextParser().parse(make_source(".txt"), io.BytesIO(b"more than eight characters"))

        with (
            mock.patch("smart_lab_index.modules.parsers.common.MAX_TABLE_ROWS", 1),
            self.assertRaisesRegex(DocumentParseError, "row limit"),
        ):
            CsvParser().parse(make_source(".csv"), make_csv())

        with (
            mock.patch(
                "smart_lab_index.modules.parsers.common.MAX_OFFICE_EXPANDED_BYTES",
                16,
            ),
            self.assertRaisesRegex(DocumentParseError, "expanded archive limit"),
        ):
            DocxParser().parse(make_source(".docx"), make_docx())

    def test_synthetic_fixture_generation_is_byte_deterministic(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            first = Path(temporary) / "first"
            second = Path(temporary) / "second"
            generate(first)
            generate(second)

            def snapshot(root: Path) -> dict[str, bytes]:
                return {
                    path.relative_to(root).as_posix(): path.read_bytes()
                    for path in sorted(root.rglob("*"))
                    if path.is_file()
                }

            self.assertEqual(snapshot(first), snapshot(second))


if __name__ == "__main__":
    unittest.main()
