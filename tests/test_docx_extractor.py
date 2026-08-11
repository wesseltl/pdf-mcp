"""Tests for Word .docx extraction. Builds its own document so it runs anywhere."""
import os
import sys
import tempfile
import types
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from pdf_mcp import docx_extractor


def make_docx(path):
    from docx import Document
    doc = Document()
    doc.add_heading("Certificate of Analysis", level=1)
    doc.add_paragraph("Batch: LAB-2026-001")
    table = doc.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "Analyte"
    header[1].text = "Result"
    header[2].text = "Unit"
    for analyte, result, unit in [
        ("pH", "7.2", ""),
        ("Glucose", "5.1", "mmol/L"),
        ("Sodium", "140", "mmol/L"),
    ]:
        cells = table.add_row().cells
        cells[0].text = analyte
        cells[1].text = result
        cells[2].text = unit
    doc.save(path)


def make_docx_with_merged_cells(path):
    from docx import Document
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].merge(table.rows[0].cells[2]).text = "Results"
    table.rows[1].cells[0].text = "pH"
    table.rows[1].cells[1].text = "7.2"
    table.rows[1].cells[2].text = ""
    doc.save(path)


class TestDocxExtractor(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.path = os.path.join(tempfile.mkdtemp(), "coa.docx")
        make_docx(cls.path)

    def test_extract_docx_text(self):
        r = docx_extractor.extract_docx_text(self.path)
        self.assertEqual(r["n_paragraphs"], 2)
        self.assertIn("Certificate of Analysis", r["text"])
        self.assertIn("Batch: LAB-2026-001", r["text"])

    def test_extract_docx_tables(self):
        r = docx_extractor.extract_docx_tables(self.path)
        self.assertEqual(r["n_tables"], 1)
        table = r["tables"][0]
        self.assertEqual(table["index"], 0)
        self.assertEqual(table["rows"][0], ["Analyte", "Result", "Unit"])
        self.assertEqual(table["rows"][2], ["Glucose", "5.1", "mmol/L"])
        self.assertEqual(table["column_count"], 3)
        self.assertFalse(table["has_merged_cells"])

    def test_extract_docx_tables_flags_merged_cells(self):
        path = os.path.join(tempfile.mkdtemp(), "merged.docx")
        make_docx_with_merged_cells(path)
        r = docx_extractor.extract_docx_tables(path)
        table = r["tables"][0]
        self.assertEqual(table["rows"][0], ["Results", "Results", "Results"])
        self.assertTrue(table["has_merged_cells"])
        self.assertFalse(table["looks_clean"])
        self.assertTrue(any("merged cells detected" in w for w in table["warnings"]))

    def test_docx_table_to_csv(self):
        csv = docx_extractor.docx_table_to_csv(self.path)
        self.assertIn("Analyte,Result,Unit", csv)
        self.assertIn("Sodium,140,mmol/L", csv)

    def test_empty_docx_reports_empty_results(self):
        from docx import Document
        path = os.path.join(tempfile.mkdtemp(), "empty.docx")
        Document().save(path)
        self.assertTrue(docx_extractor.extract_docx_text(path)["warnings"])
        self.assertTrue(docx_extractor.extract_docx_tables(path)["warnings"])

    def test_docx_table_index_must_exist(self):
        with self.assertRaises(ValueError):
            docx_extractor.docx_table_to_csv(self.path, index=1)

    def test_mcp_docx_tools_delegate_to_extractor(self):
        class FakeFastMCP:
            def __init__(self, name):
                self.name = name

            def tool(self):
                return lambda fn: fn

        sys.modules["fastmcp"] = types.SimpleNamespace(FastMCP=FakeFastMCP)
        from pdf_mcp import server

        self.assertIn("Certificate of Analysis", server.extract_docx_text(self.path)["text"])
        self.assertEqual(server.extract_docx_tables(self.path)["n_tables"], 1)
        self.assertIn("pH,7.2,", server.docx_table_to_csv(self.path))


if __name__ == "__main__":
    unittest.main()
