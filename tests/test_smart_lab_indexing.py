from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path
from typing import BinaryIO
from unittest import mock

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from smart_lab_index.application import build_application
from smart_lab_index.core.config import RuntimePolicy
from smart_lab_index.core.domain import (
    DocumentContent,
    DocumentSource,
    EntityCandidate,
    EntityRecord,
    EntityType,
    ExtractionResult,
    IndexRunStatus,
    IssueDraft,
)
from smart_lab_index.core.modules import (
    EntityRepository,
    ExtractorModule,
    FileAccess,
    IssueRepository,
    IssueRuleModule,
    ModuleCapability,
    ModuleManifest,
    ModuleType,
    NetworkAccess,
    ParserModule,
    ResolverModule,
)
from smart_lab_index.core.storage import StorageError
from smart_lab_index.modules.connectors.filesystem import FilesystemConnector


class FailingExtractor(ExtractorModule):
    manifest = ModuleManifest(
        module_id="extractor.test_failure",
        name="Failing test extractor",
        version="1.0.0",
        module_type=ModuleType.ENTITY_EXTRACTOR,
        description="Test-only module failure boundary.",
        capabilities=(ModuleCapability("extractor.test_failure", "1.0.0"),),
        network_access=NetworkAccess.NONE,
        file_access=FileAccess.NONE,
    )

    def extract(self, _document: DocumentContent) -> ExtractionResult:
        raise RuntimeError("deliberate extractor failure")


class FailOnSecondResolver(ResolverModule):
    order = 5
    manifest = ModuleManifest(
        module_id="resolver.test_failure",
        name="Failing test resolver",
        version="1.0.0",
        module_type=ModuleType.RESOLVER,
        description="Test-only transactional failure boundary.",
        capabilities=(ModuleCapability("resolver.test_failure", "1.0.0"),),
        network_access=NetworkAccess.NONE,
        file_access=FileAccess.NONE,
    )

    def __init__(self) -> None:
        super().__init__()
        self.calls = 0

    def resolve(
        self,
        _candidate: EntityCandidate,
        repository: EntityRepository,
    ) -> EntityRecord | None:
        if hasattr(repository, "connection") or hasattr(repository, "create_entity"):
            raise AssertionError("resolver received a write-capable repository")
        self.calls += 1
        if self.calls == 2:
            raise RuntimeError("deliberate resolver failure")
        return None


class CountingExtractor(ExtractorModule):
    manifest = ModuleManifest(
        module_id="extractor.test_counting",
        name="Counting test extractor",
        version="1.0.0",
        module_type=ModuleType.ENTITY_EXTRACTOR,
        description="Test-only processing-ledger probe.",
        capabilities=(ModuleCapability("extractor.test_counting", "1.0.0"),),
        configuration_schema={
            "type": "object",
            "required": ["mode"],
            "additionalProperties": False,
            "properties": {"mode": {"type": "string"}},
        },
        network_access=NetworkAccess.NONE,
        file_access=FileAccess.NONE,
    )

    def __init__(self) -> None:
        super().__init__({"mode": "first"})
        self.calls = 0

    def extract(self, _document: DocumentContent) -> ExtractionResult:
        self.calls += 1
        return ExtractionResult()


class ReplacementTextParser(ParserModule):
    priority = 10
    manifest = ModuleManifest(
        module_id="parser.test_text",
        name="Replacement text parser",
        version="1.0.0",
        module_type=ModuleType.PARSER,
        description="Test-only parser replacement.",
        capabilities=(ModuleCapability("parser.document", "1.0.0"),),
        network_access=NetworkAccess.NONE,
        file_access=FileAccess.NONE,
    )

    def supports(self, source: DocumentSource) -> bool:
        return source.content_type == "text/plain"

    def parse(self, source: DocumentSource, content: BinaryIO) -> DocumentContent:
        content.read()
        return DocumentContent(
            source_external_id=source.external_id,
            content_type=source.content_type,
            parser_module_id=self.manifest.module_id,
            parser_version=self.manifest.version,
            metadata={"replacement": True},
        )


class RepositoryProbeIssueRule(IssueRuleModule):
    manifest = ModuleManifest(
        module_id="issue.test_repository",
        name="Repository facade probe",
        version="1.0.0",
        module_type=ModuleType.ISSUE_RULE,
        description="Test-only read repository probe.",
        capabilities=(ModuleCapability("issue.test_repository", "1.0.0"),),
        network_access=NetworkAccess.NONE,
        file_access=FileAccess.NONE,
    )

    def __init__(self) -> None:
        super().__init__()
        self.received_read_only = False

    def evaluate(self, repository: IssueRepository) -> tuple[IssueDraft, ...]:
        self.received_read_only = not (
            hasattr(repository, "connection") or hasattr(repository, "create_issue")
        )
        return ()


class SmartLabIndexingTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.addCleanup(self.temporary_directory.cleanup)
        self.root = Path(self.temporary_directory.name, "sample_lab")
        self.root.mkdir()
        self.database = Path(self.temporary_directory.name, "smart-lab-index.db")

    def test_synthetic_lab_retains_conflicting_assertions_with_provenance(self):
        self._write_complete_example()

        with build_application(self.root, database=self.database) as application:
            first = application.indexing.run(application.source)

            self.assertEqual(first.status, IndexRunStatus.COMPLETED)
            self.assertEqual(first.stats["discovered"], 4)
            self.assertEqual(first.stats["parsed"], 4)
            self.assertEqual(first.stats["failed"], 0)
            self.assertEqual(len(application.store.list_entities(EntityType.LOCATION)), 2)
            self.assertEqual(len(application.store.list_entities(EntityType.ASSET)), 1)
            self.assertEqual(len(application.store.list_entities(EntityType.PERSON)), 1)

            located_in = application.store.list_active_assertions("located_in")
            self.assertEqual(len(located_in), 2)
            self.assertTrue(all(item.document_id for item in located_in))
            self.assertEqual(
                {item["parser_module_id"] for item in application.store.list_documents()},
                {"parser.docx", "parser.xlsx"},
            )
            self.assertEqual(
                {item["external_id"] for item in application.store.list_sources()},
                {
                    "equipment.xlsx",
                    "locations.xlsx",
                    "responsibilities.xlsx",
                    "SOPs/SOP_freezers.docx",
                },
            )
            self.assertEqual(
                {item.provenance["source_external_id"] for item in located_in},
                {"equipment.xlsx", "SOPs/SOP_freezers.docx"},
            )
            spreadsheet_assertion = next(
                item
                for item in located_in
                if item.provenance["source_external_id"] == "equipment.xlsx"
            )
            self.assertEqual(spreadsheet_assertion.provenance["locator"]["cell"], "D2")
            responsibilities = application.store.list_active_assertions("responsible_for")
            self.assertEqual(len(responsibilities), 1)
            self.assertEqual(
                responsibilities[0].provenance["locator"]["sheet"],
                "Responsibilities",
            )

            issues = application.store.list_issues(status="OPEN")
            self.assertEqual([issue["code"] for issue in issues], ["CONFLICTING_LOCATION"])
            observed = issues[0]["evidence"]["observed_locations"]
            self.assertEqual(
                {value["location_name"] for value in observed},
                {"Room A-101", "Room A-102"},
            )

            second = application.indexing.run(application.source)
            self.assertEqual(second.status, IndexRunStatus.COMPLETED)
            self.assertEqual(second.stats["unchanged"], 4)
            self.assertEqual(second.stats["parsed"], 0)
            self.assertEqual(len(application.store.list_assertions()), 3)
            self.assertEqual(len(application.store.list_issues()), 1)

    def test_changed_and_deleted_sources_supersede_but_do_not_destroy_history(self):
        self._write_equipment("Room A-101")
        self._write_sop("Room A-102")
        sop_path = self.root / "SOPs" / "SOP_freezers.docx"

        with build_application(self.root, database=self.database) as application:
            application.indexing.run(application.source)
            self.assertEqual(len(application.store.list_issues(status="OPEN")), 1)

            self._write_sop("Room A-101")
            changed = application.indexing.run(application.source)
            self.assertEqual(changed.stats["changed"], 1)
            self.assertEqual(changed.stats["parsed"], 1)
            self.assertEqual(len(application.store.list_active_assertions("located_in")), 2)
            self.assertEqual(len(application.store.list_assertions()), 3)
            self.assertEqual(len(application.store.list_issues(status="OPEN")), 0)
            self.assertEqual(
                application.store.list_issues(status="RESOLVED")[0]["code"],
                "CONFLICTING_LOCATION",
            )

            sop_path.unlink()
            deleted = application.indexing.run(application.source)
            self.assertEqual(deleted.stats["deleted"], 1)
            self.assertEqual(len(application.store.list_active_assertions("located_in")), 1)
            self.assertEqual(len(application.store.list_assertions()), 3)

    def test_parser_failure_isolated_from_other_documents(self):
        (self.root / "broken.pdf").write_bytes(b"not a PDF")
        (self.root / "observation.txt").write_text(
            "Freezer-001 located in Room A-101.",
            encoding="utf-8",
        )

        with build_application(self.root, database=self.database) as application:
            result = application.indexing.run(application.source)

            self.assertEqual(result.status, IndexRunStatus.COMPLETED_WITH_ERRORS)
            self.assertEqual(result.stats["discovered"], 2)
            self.assertEqual(result.stats["parsed"], 1)
            self.assertEqual(result.stats["failed"], 1)
            self.assertEqual(len(application.store.list_active_assertions("located_in")), 1)
            self.assertEqual(
                [issue["code"] for issue in application.store.list_issues(status="OPEN")],
                ["PARSING_FAILURE"],
            )

            pdf_bytes = io.BytesIO()
            document = canvas.Canvas(pdf_bytes)
            document.drawString(72, 720, "Synthetic report")
            document.save()
            (self.root / "broken.pdf").write_bytes(pdf_bytes.getvalue())
            recovered = application.indexing.run(application.source)
            self.assertEqual(recovered.stats["changed"], 1)
            self.assertEqual(recovered.stats["failed"], 0)
            self.assertEqual(application.store.list_issues(status="OPEN"), [])
            self.assertEqual(
                application.store.list_issues(status="RESOLVED")[0]["code"],
                "PARSING_FAILURE",
            )

    def test_core_state_cannot_be_created_inside_read_only_source_root(self):
        with self.assertRaisesRegex(ValueError, "outside the read-only source root"):
            build_application(self.root, database=self.root / "index.db")

    def test_optional_extractor_failure_does_not_stop_remaining_modules(self):
        (self.root / "observation.txt").write_text(
            "Freezer-001 located in Room A-101.",
            encoding="utf-8",
        )
        with build_application(self.root, database=self.database) as application:
            application.registry.register(FailingExtractor())
            self.assertEqual(application.registry.start_all(), {})

            result = application.indexing.run(application.source)

            self.assertEqual(result.status, IndexRunStatus.COMPLETED_WITH_ERRORS)
            self.assertEqual(result.stats["module_failures"], 1)
            self.assertEqual(len(application.store.list_active_assertions("located_in")), 1)
            self.assertEqual(
                [issue["code"] for issue in application.store.list_issues(status="OPEN")],
                ["MODULE_FAILURE"],
            )

    def test_failed_changed_generation_keeps_last_successful_assertions_active(self):
        path = self.root / "observation.txt"
        path.write_text("Freezer-001 located in Room A-101.", encoding="utf-8")
        with build_application(self.root, database=self.database) as application:
            application.indexing.run(application.source)
            original = application.store.list_active_assertions("located_in")[0]

            path.write_bytes(b"\xff\xfeinvalid UTF-8")
            failed = application.indexing.run(application.source)

            self.assertEqual(failed.status, IndexRunStatus.COMPLETED_WITH_ERRORS)
            active = application.store.list_active_assertions("located_in")
            self.assertEqual([item.assertion_id for item in active], [original.assertion_id])
            self.assertEqual(active[0].source_generation, 1)
            self.assertEqual(application.store.list_sources()[0]["source_generation"], 2)

    def test_deleted_source_restored_with_same_bytes_gets_a_new_active_generation(self):
        path = self.root / "observation.txt"
        content = "Freezer-001 located in Room A-101."
        path.write_text(content, encoding="utf-8")
        with build_application(self.root, database=self.database) as application:
            application.indexing.run(application.source)
            first = application.store.list_active_assertions("located_in")[0]
            path.unlink()
            application.indexing.run(application.source)
            self.assertEqual(application.store.list_active_assertions("located_in"), [])

            path.write_text(content, encoding="utf-8")
            restored = application.indexing.run(application.source)
            active = application.store.list_active_assertions("located_in")

            self.assertEqual(restored.stats["new"], 1)
            self.assertEqual(len(active), 1)
            self.assertNotEqual(active[0].assertion_id, first.assertion_id)
            self.assertEqual(active[0].source_generation, 2)
            self.assertEqual(len(application.store.list_assertions()), 2)

    def test_distinct_explicit_identifiers_are_not_merged_by_shared_name(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["asset_id", "name", "asset_type"])
        sheet.append(["ASSET-001", "Shared Name", "FREEZER"])
        sheet.append(["ASSET-002", "Shared Name", "FREEZER"])
        workbook.save(self.root / "assets.xlsx")

        with build_application(self.root, database=self.database) as application:
            application.indexing.run(application.source)
            assets = application.store.list_entities(EntityType.ASSET)

        self.assertEqual({item.identifier for item in assets}, {"ASSET-001", "ASSET-002"})

    def test_resolver_failure_rolls_back_partial_extractor_writes(self):
        self._write_equipment("Room A-101")
        resolver = FailOnSecondResolver()
        with build_application(self.root, database=self.database) as application:
            application.registry.register(resolver)
            application.registry.start_all()
            result = application.indexing.run(application.source)

            self.assertEqual(result.status, IndexRunStatus.COMPLETED_WITH_ERRORS)
            self.assertEqual(application.store.list_entities(), [])
            self.assertEqual(application.store.list_assertions(), [])
            self.assertEqual(
                application.store.list_issues(status="OPEN")[0]["code"],
                "MODULE_FAILURE",
            )

    def test_processing_ledger_retries_only_when_module_configuration_changes(self):
        (self.root / "note.txt").write_text("No relationships here.", encoding="utf-8")
        extractor = CountingExtractor()
        with build_application(self.root, database=self.database) as application:
            application.registry.register(extractor)
            application.registry.start_all()
            application.indexing.run(application.source)
            application.indexing.run(application.source)
            self.assertEqual(extractor.calls, 1)

            extractor.configuration["mode"] = "second"
            application.indexing.run(application.source)
            self.assertEqual(extractor.calls, 2)

    def test_one_connector_provider_supports_multiple_bound_source_instances(self):
        first_path = self.root / "first.txt"
        first_path.write_text("First source.", encoding="utf-8")
        second_root = Path(self.temporary_directory.name, "second_lab")
        second_root.mkdir()
        (second_root / "second.txt").write_text("Second source.", encoding="utf-8")

        with build_application(self.root, database=self.database) as application:
            application.indexing.run(application.source)
            connector = application.registry.get("connector.filesystem")
            self.assertIsInstance(connector, FilesystemConnector)
            second_source = connector.source(second_root)
            application.indexing.run(second_source)

            self.assertEqual(
                {item["source_id"] for item in application.store.list_sources()},
                {application.source.source_id, second_source.source_id},
            )
            conflicting = connector.source(
                second_root,
                source_id=application.source.source_id,
            )
            with self.assertRaisesRegex(StorageError, "different source"):
                application.indexing.run(conflicting)

    def test_higher_priority_parser_replaces_builtin_through_contract(self):
        (self.root / "note.txt").write_text("Synthetic note.", encoding="utf-8")
        with build_application(self.root, database=self.database) as application:
            application.registry.register(ReplacementTextParser())
            application.registry.start_all()
            result = application.indexing.run(application.source)

            self.assertEqual(result.status, IndexRunStatus.COMPLETED)
            self.assertEqual(
                application.store.list_documents()[0]["parser_module_id"],
                "parser.test_text",
            )

    def test_issue_modules_receive_read_only_repository_facade(self):
        (self.root / "note.txt").write_text("Synthetic note.", encoding="utf-8")
        probe = RepositoryProbeIssueRule()
        with build_application(self.root, database=self.database) as application:
            application.registry.register(probe)
            application.registry.start_all()
            application.indexing.run(application.source)

        self.assertTrue(probe.received_read_only)

    def test_failed_symlink_path_is_not_inferred_as_deleted(self):
        path = self.root / "observation.txt"
        path.write_text("Freezer-001 located in Room A-101.", encoding="utf-8")
        outside = Path(self.temporary_directory.name, "outside.txt")
        outside.write_text("outside", encoding="utf-8")
        with build_application(self.root, database=self.database) as application:
            application.indexing.run(application.source)
            path.unlink()
            path.symlink_to(outside)

            result = application.indexing.run(application.source)

            self.assertEqual(result.stats["deleted"], 0)
            self.assertEqual(len(application.store.list_active_assertions("located_in")), 1)
            self.assertIsNone(application.store.list_sources()[0]["deleted_at"])

    def test_no_egress_builtin_indexing_attempts_no_socket_connections(self):
        (self.root / "note.txt").write_text("Synthetic note.", encoding="utf-8")
        with mock.patch(
            "socket.socket.connect",
            side_effect=AssertionError("unexpected outbound connection"),
        ), build_application(
            self.root,
            database=self.database,
            policy=RuntimePolicy(no_egress=True),
        ) as application:
            result = application.indexing.run(application.source)

        self.assertEqual(result.status, IndexRunStatus.COMPLETED)

    def _write_complete_example(self) -> None:
        self._write_locations()
        self._write_equipment("Room A-101")
        self._write_responsibilities()
        self._write_sop("Room A-102")

    def _write_locations(self) -> None:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Locations"
        sheet.append(["location_id", "name", "location_type"])
        sheet.append(["Room A-101", "Room A-101", "ROOM"])
        sheet.append(["Room A-102", "Room A-102", "ROOM"])
        workbook.save(self.root / "locations.xlsx")

    def _write_equipment(self, location: str) -> None:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Assets"
        sheet.append(["asset_id", "name", "asset_type", "location"])
        sheet.append(["Freezer-001", "Freezer-001", "FREEZER", location])
        workbook.save(self.root / "equipment.xlsx")

    def _write_responsibilities(self) -> None:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Responsibilities"
        sheet.append(["person", "asset_id", "relationship"])
        sheet.append(["Alex Example", "Freezer-001", "responsible_for"])
        workbook.save(self.root / "responsibilities.xlsx")

    def _write_sop(self, location: str) -> None:
        directory = self.root / "SOPs"
        directory.mkdir(exist_ok=True)
        document = Document()
        document.add_heading("Freezer procedure", level=1)
        document.add_paragraph(f"Freezer-001 located in {location}.")
        document.save(directory / "SOP_freezers.docx")


if __name__ == "__main__":
    unittest.main()
